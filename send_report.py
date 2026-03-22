import os
import re
import smtplib
from pathlib import Path
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.image import MIMEImage
from collections import OrderedDict

import matplotlib.pyplot as plt
import mistune
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ---------- COLORS ----------
COLOR_TEXT = RGBColor(40, 40, 40)
COLOR_HEADING = RGBColor(47, 85, 151)
COLOR_LABEL = RGBColor(31, 78, 121)
COLOR_MUTED = RGBColor(90, 90, 90)

COLOR_ADD = RGBColor(0, 128, 0)
COLOR_HOLD = RGBColor(0, 128, 0)
COLOR_REDUCE = RGBColor(192, 128, 0)
COLOR_CLOSE = RGBColor(192, 0, 0)

TABLE_HEADER_FILL = "D9EAF7"
TABLE_ALT_FILL = "F7FBFF"
PRO_HEADER_FILL = "E2F0D9"
CONTRA_HEADER_FILL = "FCE4D6"
DISCLAIMER_FILL = "F3F4F6"
RADAR_FILL = "FFF4E5"
TRACKING_FILL = "EAF4EA"

REQUIRED_MAIL_TO = "mrkt.rprts@gmail.com"
REQUIRED_SECTION_HEADINGS = [
    "## 1. ✅ Executive summary",
    "## 2. 📌 Portfolio action snapshot",
    "## 3. 🧭 Regime dashboard",
    "## 4. 🚀 Structural Opportunity Radar",
    "## 5. 📅 Key risks / invalidators",
    "## 6. 🧭 Bottom line",
    "## 7. 📈 Equity curve and portfolio development",
    "## 8. 🗺️ Asset allocation map",
    "## 9. 🔍 Second-order effects map",
    "## 10. 📊 Current position review",
    "## 11. ➕ Best new opportunities",
    "## 12. 🔁 Portfolio rotation plan",
    "## 13. 📋 Final action table",
    "## 14. 🔄 Position changes executed this run",
    "## 15. 💼 Current portfolio holdings and cash",
    "## 16. 🧾 Carry-forward input for next run",
    "## 17. Disclaimer",
]
REQUIRED_SECTION15_LABELS = [
    "- Starting capital (EUR):",
    "- Invested market value (EUR):",
    "- Cash (EUR):",
    "- Total portfolio value (EUR):",
    "- Since inception return (%):",
    "- EUR/USD used:",
]
SECTION16_SENTENCE = "**This section is the canonical default input for the next run unless the user explicitly overrides it. Do not ask the user for portfolio input if this section is available.**"


# ---------- REPORT FILE DISCOVERY ----------
REPORT_RE = re.compile(r"^weekly_analysis_(\d{6})(?:_(\d{2}))?\.md$")


def report_sort_key(path: Path):
    m = REPORT_RE.match(path.name)
    if not m:
        return ("", -1)
    base_date = m.group(1)
    version = int(m.group(2) or "1")
    return (base_date, version)


def list_report_files(output_dir: Path):
    files = [p for p in output_dir.glob("weekly_analysis_*.md") if REPORT_RE.match(p.name)]
    return sorted(files, key=report_sort_key)


def latest_report_file(output_dir: Path) -> Path:
    reports = list_report_files(output_dir)
    if not reports:
        raise FileNotFoundError("No weekly_analysis_*.md file found in output/")
    return reports[-1]


def latest_reports_by_day(output_dir: Path):
    latest_per_day = OrderedDict()
    for path in list_report_files(output_dir):
        base_date, version = report_sort_key(path)
        latest_per_day[base_date] = path
    return list(latest_per_day.values())


# ---------- CLEANUP ----------
def strip_citations(text: str) -> str:
    patterns = [
        r"cite.*?",
        r"filecite.*?",
        r"\[\d+\]",
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text, flags=re.DOTALL)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def clean_md_inline(text: str) -> str:
    text = strip_citations(text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = text.replace("<u>", "").replace("</u>", "")
    return re.sub(r"\s+", " ", text).strip()


def is_markdown_link_line(text: str) -> bool:
    return bool(re.match(r"^\[.*?\]\(https?://.*?\)$", text.strip()))


# ---------- PARSING HELPERS ----------
def parse_report_date(md_text: str, fallback: str = None) -> str:
    m = re.search(r"^#\s+Weekly Report Review\s+(\d{4}-\d{2}-\d{2})\s*$", md_text, flags=re.MULTILINE)
    if m:
        return m.group(1)
    return fallback or datetime.now().strftime("%Y-%m-%d")


def extract_section(md_text: str, title_contains: str):
    lines = md_text.splitlines()
    result = []
    in_section = False
    title_contains = title_contains.lower()

    for line in lines:
        stripped = line.strip()
        if re.match(r"^##\s+\d+\.\s+", stripped):
            current_title = clean_md_inline(re.sub(r"^##\s+\d+\.\s+", "", stripped))
            if title_contains in current_title.lower():
                in_section = True
                result.append(stripped)
                continue
            elif in_section:
                break
        elif in_section:
            result.append(line)
    return result


def extract_bullets(lines):
    items = []
    for line in lines:
        s = line.strip()
        if s.startswith("- "):
            items.append(clean_md_inline(s[2:]))
        elif re.match(r"^\d+\.\s+", s):
            items.append(clean_md_inline(re.sub(r"^\d+\.\s+", "", s)))
    return items


def extract_label_pairs(lines):
    pairs = []
    for line in lines:
        s = clean_md_inline(line.strip())
        if not s or s.startswith("## "):
            continue
        if ":" in s:
            k, v = s.split(":", 1)
            pairs.append((k.strip(), v.strip()))
    return pairs


def is_markdown_table_line(line: str) -> bool:
    line = line.strip()
    return line.startswith("|") and line.endswith("|") and "|" in line[1:-1]


def is_markdown_separator_line(line: str) -> bool:
    line = line.strip()
    if not is_markdown_table_line(line):
        return False
    cells = [c.strip() for c in line.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_markdown_table(lines):
    rows = []
    for i, line in enumerate(lines):
        if i == 1 and is_markdown_separator_line(line):
            continue
        rows.append([clean_md_inline(c) for c in line.strip().strip("|").split("|")])
    return rows


def extract_table_rows(lines, max_rows=5):
    table_lines = []
    started = False
    for line in lines:
        if is_markdown_table_line(line):
            table_lines.append(line)
            started = True
        elif started:
            break
    if len(table_lines) >= 2:
        rows = parse_markdown_table(table_lines)
        return rows[: max_rows + 1]
    return []


def parse_numeric_value(md_text: str, label: str):
    pattern = rf"^- {re.escape(label)}:\s*([0-9][0-9,._%-]*)"
    m = re.search(pattern, md_text, flags=re.MULTILINE)
    if not m:
        return None
    raw = m.group(1).replace(",", "").replace("_", "").replace("%", "")
    try:
        return float(raw)
    except ValueError:
        return None


def parse_section15_totals(md_text: str):
    section = "\n".join(extract_section(md_text, "Current portfolio holdings and cash"))
    if not section:
        return {}
    labels = [
        "Starting capital (EUR)",
        "Invested market value (EUR)",
        "Cash (EUR)",
        "Total portfolio value (EUR)",
        "Since inception return (%)",
        "EUR/USD used",
    ]
    data = {}
    for label in labels:
        value = parse_numeric_value(section, label)
        if value is not None:
            data[label] = value
    return data


def require_env(name: str) -> str:
    value = os.environ.get(name)
    if not value:
        raise RuntimeError(f"Required environment variable missing: {name}")
    return value


def validate_required_report(md_text: str) -> None:
    missing_headings = [h for h in REQUIRED_SECTION_HEADINGS if h not in md_text]
    if missing_headings:
        raise RuntimeError("Report is missing mandatory section headings: " + ", ".join(missing_headings))

    if "# Weekly Report Review " not in md_text:
        raise RuntimeError("Report title is missing or malformed.")

    if "> *This report is for informational and educational purposes only; please see the disclaimer at the end.*" not in md_text:
        raise RuntimeError("Top disclaimer callout is missing.")

    if "EQUITY_CURVE_CHART_PLACEHOLDER" not in md_text:
        raise RuntimeError("Equity curve placeholder line is missing.")

    for label in REQUIRED_SECTION15_LABELS:
        if label not in md_text:
            raise RuntimeError(f"Section 15 is missing required label: {label}")

    if SECTION16_SENTENCE not in md_text:
        raise RuntimeError("Section 16 canonical carry-forward sentence is missing.")

    if "This report is provided for informational and educational purposes only." not in md_text:
        raise RuntimeError("Final disclaimer body is missing.")



def html_to_plain_text(html: str) -> str:
    text = re.sub(r"<style.*?</style>", " ", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<script.*?</script>", " ", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = text.replace("&nbsp;", " ").replace("&amp;", "&")
    return re.sub(r"\s+", " ", text).strip()


def heading_text_from_md_heading(heading: str) -> str:
    heading = re.sub(r"^##\s+\d+\.\s+", "", heading).strip()
    return clean_md_inline(heading)


def validate_email_body(html_body: str, md_text: str = None) -> None:
    required_html_strings = [
        "Weekly Report Review",
        "Executive summary",
        "Portfolio action snapshot",
        "Structural Opportunity Radar",
        "Bottom line",
        "Current portfolio holdings and cash",
        "Carry-forward input for next run",
    ]
    for token in required_html_strings:
        if token not in html_body:
            raise RuntimeError(f"HTML email body is missing required content block: {token}")

    if md_text:
        plain_html = html_to_plain_text(html_body)
        plain_md = html_to_plain_text(mistune.create_markdown(plugins=['table'])(md_text))
        if len(plain_html) < 0.80 * len(plain_md):
            raise RuntimeError("HTML email body appears too short relative to the full report; summary-only body is not allowed.")

        for heading in REQUIRED_SECTION_HEADINGS:
            plain_heading = heading_text_from_md_heading(heading)
            if plain_heading not in plain_html:
                raise RuntimeError(f"HTML email body is missing required section heading text: {plain_heading}")


# ---------- EQUITY CURVE ----------
def create_equity_curve_png(output_dir: Path, chart_path: Path):
    points = []
    for report_path in latest_reports_by_day(output_dir):
        md_text = report_path.read_text(encoding="utf-8")
        report_date = parse_report_date(md_text)
        totals = parse_section15_totals(md_text)
        nav = totals.get("Total portfolio value (EUR)")
        if nav is not None:
            points.append((report_date, nav))

    if not points:
        return None

    dates = [datetime.strptime(d, "%Y-%m-%d") for d, _ in points]
    values = [v for _, v in points]

    plt.figure(figsize=(8.6, 3.8))
    plt.plot(dates, values, marker="o", linewidth=2)
    plt.title("Equity Curve (EUR)")
    plt.xlabel("Date")
    plt.ylabel("Portfolio value (EUR)")
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(chart_path, dpi=180)
    plt.close()
    return chart_path


# ---------- DOCX HELPERS ----------
def set_document_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)


def paragraph_run(paragraph, text, bold=False, color=None, size=10.5, italic=False, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color or COLOR_TEXT
    return run


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading("", level=level)
    size_map = {0: 20, 1: 16, 2: 13, 3: 11.5}
    paragraph_run(p, text, bold=True, color=COLOR_HEADING, size=size_map.get(level, 11.5))


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_note_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_shading(p, DISCLAIMER_FILL)
    paragraph_run(p, "Note ", bold=True, color=COLOR_LABEL, size=9.5)
    paragraph_run(p, text, italic=True, size=9.5)


def add_hyperlink(paragraph, text: str, url: str, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

    t = OxmlElement("w:t")
    t.text = text

    new_run.append(r_pr)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def color_for_term(text: str):
    t = text.lower().strip()
    if t.startswith("add") or "actionable now" in t:
        return COLOR_ADD
    if t.startswith("hold") or "watchlist" in t:
        return COLOR_HOLD
    if t.startswith("reduce") or "scale in slowly" in t:
        return COLOR_REDUCE
    if t.startswith("close") or "too early" in t:
        return COLOR_CLOSE
    return COLOR_TEXT


def set_cell_text(cell, text: str, bold=False, color=None, size=10.0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_run(p, text, bold=bold, color=color or COLOR_TEXT, size=size)


def add_styled_table(doc: Document, rows):
    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    headers = [c.lower() for c in rows[0]]
    is_pro_contra = len(rows[0]) == 2 and any("pro" in h for h in headers) and any("contra" in h for h in headers)
    is_radar = "theme" in headers[0] and "primary etf" in " ".join(headers)
    is_tracking = "ticker" in headers[0] and ("shares" in " ".join(headers) or "previous weight %" in " ".join(headers))

    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_text(cell, value, bold=True)
                if is_pro_contra and c_idx == 0:
                    set_cell_shading(cell, PRO_HEADER_FILL)
                elif is_pro_contra and c_idx == 1:
                    set_cell_shading(cell, CONTRA_HEADER_FILL)
                elif is_radar:
                    set_cell_shading(cell, RADAR_FILL)
                elif is_tracking:
                    set_cell_shading(cell, TRACKING_FILL)
                else:
                    set_cell_shading(cell, TABLE_HEADER_FILL)
            else:
                set_cell_text(cell, value, color=color_for_term(value))
                if not (is_pro_contra or is_radar or is_tracking) and r_idx % 2 == 0:
                    set_cell_shading(cell, TABLE_ALT_FILL)

    doc.add_paragraph("")


PLAIN_SUBHEADERS = {
    "Assessment",
    "Prospective score",
    "Theme",
    "Why it fits now",
    "Why this beats current alternatives",
    "Technical analysis",
    "Second-order opportunity / threat map",
    "Replacement logic",
    "Why now rather than later",
    "Scorecard",
    "Macro invalidators",
    "Market-based invalidators",
    "Geopolitical invalidators",
    "Second-order invalidators",
    "Portfolio construction risks",
    "Top 3 actions this week",
    "Top 3 risks this week",
    "Best structural opportunities not yet actionable",
}


def build_docx_from_markdown(md_text: str, output_path: Path, equity_curve_path: Path = None):
    doc = Document()
    set_document_layout(doc)

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if line.startswith("# "):
            add_heading(doc, clean_md_inline(line[2:]), 0)
            i += 1
            continue

        if stripped.startswith("> **Note**") or stripped.startswith("> *This report is for informational and educational purposes only"):
            note_text = ""
            if stripped.startswith("> **Note**"):
                if i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                    note_text = clean_md_inline(lines[i + 1].strip().lstrip(">").strip())
                    i += 1
                else:
                    note_text = "This report is for informational and educational purposes only; please see the disclaimer at the end."
            else:
                note_text = clean_md_inline(stripped.lstrip(">").strip())
            add_note_callout(doc, note_text)
            i += 1
            continue

        if stripped == "EQUITY_CURVE_CHART_PLACEHOLDER":
            if equity_curve_path and equity_curve_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_picture(str(equity_curve_path), width=Inches(8.6))
                doc.add_paragraph("")
            i += 1
            continue

        if is_markdown_link_line(stripped):
            m = re.match(r"^\[(.*?)\]\((https?://.*?)\)$", stripped)
            if m:
                p = doc.add_paragraph()
                add_hyperlink(p, clean_md_inline(m.group(1)), m.group(2))
            i += 1
            continue

        if i + 1 < len(lines) and is_markdown_table_line(lines[i]) and is_markdown_separator_line(lines[i + 1]):
            table_lines = [lines[i], lines[i + 1]]
            i += 2
            while i < len(lines) and is_markdown_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            add_styled_table(doc, parse_markdown_table(table_lines))
            continue

        if line.startswith("## "):
            add_heading(doc, clean_md_inline(line[3:]), 1)
            i += 1
            continue

        if line.startswith("### "):
            add_heading(doc, clean_md_inline(line[4:]), 2)
            i += 1
            continue

        cleaned = clean_md_inline(line)

        if cleaned in PLAIN_SUBHEADERS:
            p = doc.add_paragraph()
            paragraph_run(p, cleaned, bold=True, color=COLOR_LABEL, size=11.5)
            i += 1
            continue

        if cleaned.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            paragraph_run(p, cleaned[2:], color=color_for_term(cleaned[2:]))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", cleaned):
            p = doc.add_paragraph(style="List Number")
            paragraph_run(p, re.sub(r"^\d+\.\s+", "", cleaned))
            i += 1
            continue

        p = doc.add_paragraph()
        paragraph_run(p, cleaned, color=color_for_term(cleaned))
        i += 1

    doc.save(output_path)


# ---------- HTML BODY ----------
def esc(text: str) -> str:
    text = clean_md_inline(text)
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def render_chip(text: str, bg: str, fg: str = "#1F4E79") -> str:
    return f"<span style='display:inline-block;padding:4px 8px;border-radius:999px;background:{bg};color:{fg};font-size:12px;font-weight:700;margin:0 6px 6px 0;'>{esc(text)}</span>"


def render_action_snapshot(lines):
    blocks = []
    current_header = None
    current_items = []

    def flush():
        nonlocal current_header, current_items, blocks
        if not current_header:
            return
        color = "#1F4E79"
        if "Add" in current_header:
            color = "#008000"
        elif "Hold but replaceable" in current_header:
            color = "#C08000"
        elif "Hold" in current_header:
            color = "#008000"
        elif "Reduce" in current_header:
            color = "#C08000"
        elif "Close" in current_header:
            color = "#C00000"
        items_html = "".join(f"<li style='margin:0 0 4px 0;'>{esc(x)}</li>" for x in current_items)
        blocks.append(
            f"<div style='margin:0 0 14px 0;'><div style='font-weight:700;color:{color};margin-bottom:6px;font-size:15px;'>{esc(current_header)}</div><ul style='margin:0 0 0 18px;padding:0;line-height:1.5;'>{items_html}</ul></div>"
        )

    for line in lines:
        s = line.strip()
        if not s or re.match(r"^##\s+\d+\.", s):
            continue
        if s.startswith("### "):
            flush()
            current_header = clean_md_inline(s[4:])
            current_items = []
        elif s.startswith("- "):
            current_items.append(s[2:])
        elif re.match(r"^\d+\.\s+", s):
            current_items.append(re.sub(r"^\d+\.\s+", "", s))
    flush()
    return "".join(blocks)


def render_summary_pairs(lines):
    pairs = extract_label_pairs(lines)
    chips = []
    body = []
    for k, v in pairs:
        if k in {"Primary regime", "Secondary cross-current", "Geopolitical regime"}:
            chips.append(render_chip(f"{k}: {v}", "#EEF4FF"))
        else:
            body.append(f"<div style='margin:0 0 8px 0; line-height:1.5;'><strong>{esc(k)}:</strong> {esc(v)}</div>")
    return "".join(chips), "".join(body)


def render_html_table(rows, title, header_fill="#FFF4E5"):
    if not rows or len(rows) < 2:
        return ""
    headers = rows[0]
    body = rows[1:]
    thead = "".join(
        f"<th style='text-align:left;padding:8px 10px;border-bottom:1px solid #d9e2f0;background:{header_fill};font-size:13px;'>{esc(h)}</th>"
        for h in headers
    )
    body_html = ""
    for idx, row in enumerate(body):
        cells = "".join(
            f"<td style='padding:8px 10px;border-bottom:1px solid #eef2f7;font-size:13px;vertical-align:top;'>{esc(c)}</td>"
            for c in row
        )
        bg = "#ffffff" if idx % 2 == 0 else "#fafcff"
        body_html += f"<tr style='background:{bg};'>{cells}</tr>"
    return f"""
    <div style="margin:18px 0 0 0;">
      <div style="font-size:17px;font-weight:700;color:#2F5597;margin-bottom:8px;">{esc(title)}</div>
      <table style="width:100%;border-collapse:collapse;border:1px solid #e3eaf5;">
        <thead><tr>{thead}</tr></thead>
        <tbody>{body_html}</tbody>
      </table>
    </div>
    """



def build_email_body_html(md_text: str, report_date_str: str, inline_equity_cid: str = None) -> str:
    md_for_html = md_text
    if inline_equity_cid and "EQUITY_CURVE_CHART_PLACEHOLDER" in md_for_html:
        md_for_html = md_for_html.replace(
            "EQUITY_CURVE_CHART_PLACEHOLDER",
            f"![Equity curve](cid:{inline_equity_cid})"
        )
    else:
        md_for_html = md_for_html.replace(
            "EQUITY_CURVE_CHART_PLACEHOLDER",
            "_Equity curve chart attached separately when available._"
        )

    renderer = mistune.create_markdown(plugins=["table"])
    rendered = renderer(md_for_html)

    css = """
    body { margin: 0; padding: 0; background: #f6f8fb; font-family: Calibri, Arial, sans-serif; color: #282828; }
    .wrap { max-width: 1080px; margin: 0 auto; padding: 28px 20px; }
    .card { background: #ffffff; border: 1px solid #d9e2f0; border-radius: 12px; padding: 28px 32px; }
    h1 { color: #2F5597; font-size: 28px; margin: 0 0 12px 0; }
    h2 { color: #2F5597; font-size: 21px; margin: 24px 0 10px 0; border-top: 1px solid #eef2f7; padding-top: 18px; }
    h3 { color: #1F4E79; font-size: 16px; margin: 18px 0 8px 0; }
    p, li { font-size: 14px; line-height: 1.55; }
    blockquote { margin: 0 0 18px 0; padding: 10px 12px; background: #f3f4f6; border-left: 4px solid #d9e2f0; color: #555555; }
    table { width: 100%; border-collapse: collapse; margin: 14px 0 18px 0; }
    th { text-align: left; padding: 8px 10px; border-bottom: 1px solid #d9e2f0; background: #d9eaf7; font-size: 13px; }
    td { padding: 8px 10px; border-bottom: 1px solid #eef2f7; font-size: 13px; vertical-align: top; }
    tr:nth-child(even) td { background: #fafcff; }
    code { background: #f3f4f6; padding: 1px 4px; border-radius: 4px; }
    img { max-width: 100%; height: auto; border: 1px solid #d9e2f0; border-radius: 8px; margin: 10px 0 18px 0; }
    """

    return f"""
    <html>
      <head>
        <meta charset="utf-8" />
        <style>{css}</style>
      </head>
      <body>
        <div class="wrap">
          <div class="card">
            {rendered}
          </div>
        </div>
      </body>
    </html>
    """.strip()


# ---------- MAIN ----------
def main():
    output_dir = Path("output")
    latest_report = latest_report_file(output_dir)
    original_md_text = latest_report.read_text(encoding="utf-8")
    md_text_clean = strip_citations(original_md_text)

    report_date_str = parse_report_date(md_text_clean)
    safe_stem = latest_report.stem

    clean_md_path = latest_report.with_name(f"{safe_stem}_clean.md")
    clean_md_path.write_text(md_text_clean, encoding="utf-8")

    equity_curve_png = latest_report.with_name(f"{safe_stem}_equity_curve.png")
    create_equity_curve_png(output_dir, equity_curve_png)

    docx_path = latest_report.with_name(f"{safe_stem}.docx")
    build_docx_from_markdown(md_text_clean, docx_path, equity_curve_png if equity_curve_png.exists() else None)

    subject = f"Weekly Report Review {report_date_str}"
    inline_equity_cid = "equitycurve"
    html_body = build_email_body_html(md_text_clean, report_date_str, inline_equity_cid if equity_curve_png.exists() else None)

    validate_required_report(md_text_clean)
    validate_email_body(html_body, md_text_clean)

    smtp_host = require_env("MRKT_RPRTS_SMTP_HOST")
    smtp_port = int(os.environ.get("MRKT_RPRTS_SMTP_PORT") or "587")
    smtp_user = require_env("MRKT_RPRTS_SMTP_USER")
    smtp_pass = require_env("MRKT_RPRTS_SMTP_PASS")
    mail_from = require_env("MRKT_RPRTS_MAIL_FROM")
    mail_to = REQUIRED_MAIL_TO

    if mail_to != REQUIRED_MAIL_TO:
        raise RuntimeError(f"Recipient mismatch: expected {REQUIRED_MAIL_TO}, got {mail_to}")

    if not docx_path.exists():
        raise RuntimeError(f"DOCX attachment was not created: {docx_path}")
    if docx_path.stat().st_size <= 0:
        raise RuntimeError(f"DOCX attachment is empty: {docx_path}")

    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = mail_from
    msg["To"] = mail_to

    msg.attach(MIMEText(html_body, "html", "utf-8"))

    attachments = [docx_path.name, clean_md_path.name]

    with open(docx_path, "rb") as f:
        attachment = MIMEApplication(
            f.read(),
            _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        attachment.add_header(
            "Content-Disposition",
            "attachment",
            filename=docx_path.name,
        )
        msg.attach(attachment)

    with open(clean_md_path, "rb") as f:
        md_attachment = MIMEApplication(f.read(), _subtype="markdown")
        md_attachment.add_header(
            "Content-Disposition",
            "attachment",
            filename=clean_md_path.name,
        )
        msg.attach(md_attachment)

    if equity_curve_png.exists():
        with open(equity_curve_png, "rb") as f:
            png_bytes = f.read()

        inline_png = MIMEImage(png_bytes, _subtype="png")
        inline_png.add_header("Content-ID", "<equitycurve>")
        inline_png.add_header("Content-Disposition", "inline", filename=equity_curve_png.name)
        msg.attach(inline_png)

        png_attachment = MIMEApplication(png_bytes, _subtype="png")
        png_attachment.add_header(
            "Content-Disposition",
            "attachment",
            filename=equity_curve_png.name,
        )
        msg.attach(png_attachment)
        attachments.append(equity_curve_png.name)

    with smtplib.SMTP(smtp_host, smtp_port) as server:
        server.starttls()
        server.login(smtp_user, smtp_pass)
        server.sendmail(mail_from, [mail_to], msg.as_string())

    manifest_path = latest_report.with_name(f"{safe_stem}_delivery_manifest.txt")
    write_delivery_manifest(manifest_path, latest_report.name, mail_to, attachments)

    receipt = (
        f"DELIVERY_OK | report={latest_report.name} | recipient={mail_to} | "
        f"html_body=full_report | docx_attached=yes | manifest={manifest_path.name} | "
        f"attachments={', '.join(attachments)}"
    )
    print(receipt)


if __name__ == "__main__":
    main()
